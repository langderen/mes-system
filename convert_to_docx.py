"""
将Markdown文档转换为docx格式
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def convert_md_to_docx(md_file, docx_file):
    """将Markdown文件转换为docx文件"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # 设置中文字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # 处理标题
        if line.startswith('# '):
            # 一级标题
            heading = doc.add_heading(level=1)
            run = heading.runs[0] if heading.runs else heading.add_run()
            run.text = line[2:].strip()
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
        elif line.startswith('## '):
            # 二级标题
            heading = doc.add_heading(level=2)
            run = heading.runs[0] if heading.runs else heading.add_run()
            run.text = line[3:].strip()
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
        elif line.startswith('### '):
            # 三级标题
            heading = doc.add_heading(level=3)
            run = heading.runs[0] if heading.runs else heading.add_run()
            run.text = line[4:].strip()
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
        elif line.startswith('#### '):
            # 四级标题
            heading = doc.add_heading(level=4)
            run = heading.runs[0] if heading.runs else heading.add_run()
            run.text = line[5:].strip()
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
        elif line.startswith('##### '):
            # 五级标题
            heading = doc.add_heading(level=5)
            run = heading.runs[0] if heading.runs else heading.add_run()
            run.text = line[6:].strip()
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 处理代码块
        elif line.startswith('```'):
            # 收集代码块内容
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            # 添加代码块
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                
                # 添加代码背景
                pPr = p._element.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F5F5F5')
                shd.set(qn('w:val'), 'clear')
                pPr.append(shd)
                
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        # 处理表格
        elif '|' in line and i + 1 < len(lines) and '|---' in lines[i+1]:
            # 收集表格数据
            table_data = []
            # 表头
            header = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(header)
            
            # 跳过分隔行
            i += 2
            
            # 数据行
            while i < len(lines) and '|' in lines[i] and lines[i].strip():
                row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                table_data.append(row)
                i += 1
            
            # 创建表格
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                
                # 设置表格宽度
                table.autofit = True
                
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_text in enumerate(row_data):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = cell_text
                        
                        # 设置单元格格式
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                                run.font.name = '微软雅黑'
                                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                                
                                # 表头加粗和背景色
                                if row_idx == 0:
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                    # 添加背景色
                                    shading = OxmlElement('w:shd')
                                    shading.set(qn('w:fill'), '2980B9')
                                    shading.set(qn('w:val'), 'clear')
                                    cell._tc.get_or_add_tcPr().append(shading)
            
            continue  # 跳过i的递增，因为已经在循环中处理了
        
        # 处理列表
        elif re.match(r'^\d+\.\s', line):
            # 有序列表
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            
        elif line.startswith('- ') or line.startswith('* '):
            # 无序列表
            text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
        
        # 处理空行
        elif line.strip() == '':
            doc.add_paragraph('')
        
        # 处理普通段落
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)
        
        i += 1
    
    # 保存文档
    doc.save(docx_file)
    print(f'转换完成: {docx_file}')


def add_formatted_text(paragraph, text):
    """为段落添加格式化文本（处理粗体、斜体等）"""
    # 处理粗体 **text**
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if i % 2 == 1:  # 粗体部分
            run = paragraph.add_run(part)
            run.bold = True
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        else:
            # 处理行内代码 `code`
            sub_parts = re.split(r'`(.+?)`', part)
            for j, sub_part in enumerate(sub_parts):
                if j % 2 == 1:  # 代码部分
                    run = paragraph.add_run(sub_part)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
                else:
                    run = paragraph.add_run(sub_part)
                    run.font.name = '微软雅黑'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


if __name__ == '__main__':
    md_file = r'E:\mes\功能文档-详细版.md'
    docx_file = r'E:\mes\功能文档-详细版.docx'
    convert_md_to_docx(md_file, docx_file)
